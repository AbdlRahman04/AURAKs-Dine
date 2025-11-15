from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
document = Document()

# Add title
title = document.add_heading('QuickDineFlow - Functional Requirements', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add introduction paragraph
document.add_paragraph(
    'This document contains detailed functional requirements for the QuickDineFlow '
    'cafeteria ordering platform. Each requirement is presented in a structured table format.'
)
document.add_paragraph()

# Define all functional requirements
requirements = [
    {
        "title": "Order Payment Processing",
        "rows": {
            "Function": "Process online and cash payments for confirmed cafeteria orders.",
            "Description": ("Handles payment for orders using multiple methods, including card "
                          "payments via Stripe and cash on pickup. Ensures that payment details "
                          "are securely transmitted and that the order is only marked as paid "
                          "when the payment provider confirms success."),
            "Inputs": "Order identifier, selected payment method, payment amount, card/payment token (for card payments), confirmation of cash selection.",
            "Source": "Checkout page inputs and payment method selection by the student.",
            "Outputs": "Payment confirmation or failure result; updated order payment status.",
            "Destination": "Orders database, payment status display on the student interface, and admin dashboards.",
            "Action": ("When the student confirms payment, the system validates the order amount and payment method, "
                      "initiates the payment with the provider if card is chosen, receives the result, and updates "
                      "the order record accordingly. For cash, the system marks the order as \"payment pending\" until pickup."),
            "Requirements": "Reliable connection to Stripe; valid order in the system; secure backend for handling tokens.",
            "Pre-condition": "Order has passed validation and is ready for payment; student is authenticated.",
            "Post-condition": "Order payment status is updated to completed, pending, or failed based on the outcome.",
            "Side effects": "Payment attempts are logged for audit; failed payments may trigger user notifications."
        }
    },
    {
        "title": "Payment Status and Digital Receipts",
        "rows": {
            "Function": "Generate and display payment status and digital receipts for student orders.",
            "Description": ("Provides students with clear confirmation of successful payments, including order details, "
                          "amount paid, and payment method. Receipts can be viewed later from the order history."),
            "Inputs": "Order data, payment result, timestamp, payment method details.",
            "Source": "Orders database and payment processing module.",
            "Outputs": "Receipt record associated with the order and a rendered receipt view.",
            "Destination": "Student order confirmation page, order history page, and backend receipts store.",
            "Action": ("After a successful payment, the system compiles order and payment information into a receipt object, "
                      "stores it, and displays a confirmation screen with a unique order number and payment summary."),
            "Requirements": "Consistent order and payment schemas; storage for historical receipts.",
            "Pre-condition": "Payment for an order has succeeded.",
            "Post-condition": "A persistent digital receipt exists and is accessible to the student.",
            "Side effects": "Receipt data may be included in analytics and export reports."
        }
    },
    {
        "title": "Saved Payment Methods",
        "rows": {
            "Function": "Manage secure storage and reuse of student payment methods.",
            "Description": ("Allows students to save card information (via Stripe payment method IDs) for faster checkout "
                          "on future orders, while ensuring that sensitive card data is never stored directly on the café system."),
            "Inputs": "Stripe payment method identifier, student identifier, card label (e.g., \"Personal Visa\").",
            "Source": "Payment form during checkout or dedicated \"Manage Payment Methods\" screen.",
            "Outputs": "Updated payment-method records linked to the student account.",
            "Destination": "Payment methods table in the backend; selection list on the checkout page.",
            "Action": ("When the student chooses to save or delete a payment method, the system communicates with Stripe "
                      "to attach/detach the method to the customer and updates local records, marking a default card when appropriate."),
            "Requirements": "Stripe customer profile for each student; secure server-to-Stripe communication.",
            "Pre-condition": "Student is authenticated; at least one successful or valid card entry exists.",
            "Post-condition": "Payment methods list is updated; default payment method is clearly identified.",
            "Side effects": "Changes are logged in the audit trail for security compliance."
        }
    },
    {
        "title": "Pickup Time Slot Management",
        "rows": {
            "Function": "Allocate and validate pickup time slots for student orders.",
            "Description": ("Ensures that each order is assigned a valid pickup time at least 15 minutes in the future, "
                          "while preventing overbooking of specific time windows."),
            "Inputs": "Requested pickup time, current time, café configuration (slot size and capacity), existing orders.",
            "Source": "Student selection on the checkout page; order database.",
            "Outputs": "Confirmed pickup time or alternative suggested time slot.",
            "Destination": "Order record, kitchen display, and student confirmation page.",
            "Action": ("When a student selects a time, the system checks that it is at least 15 minutes ahead and that "
                      "the number of orders in that slot is below the configured capacity; if not, it proposes the nearest available slot."),
            "Requirements": "Configurable slot size and capacity; synchronized server time.",
            "Pre-condition": "Order is valid and has not yet been placed.",
            "Post-condition": "Order record stores a confirmed pickup time.",
            "Side effects": "Slot capacity statistics are updated for analytics."
        }
    },
    {
        "title": "Order Identification and Numbering",
        "rows": {
            "Function": "Generate a unique order identifier for each confirmed order.",
            "Description": ("Creates human-readable, unique order numbers used in confirmations, kitchen display, and analytics."),
            "Inputs": "Order creation event, timestamp, random or sequential component.",
            "Source": "Checkout process after all validations pass.",
            "Outputs": "Unique order number string.",
            "Destination": "Order record, student confirmation screen, receipts, and admin views.",
            "Action": ("When the order is finalized, the system generates a new identifier using a timestamp and random component, "
                      "verifies its uniqueness, then attaches it to the order."),
            "Requirements": "Collision-safe ID generation logic; access to order store to ensure uniqueness.",
            "Pre-condition": "Order has passed payment or cash-selection validation.",
            "Post-condition": "Order has a permanent unique identifier that can be used for lookup.",
            "Side effects": "Order IDs may be used as keys in logs and external integrations."
        }
    },
    {
        "title": "Order Status Tracking",
        "rows": {
            "Function": "Track and update order status throughout its lifecycle.",
            "Description": ("Represents each order using standard states – received, preparing, ready, completed, and cancelled – "
                          "and ensures that status changes are reflected consistently for students and admins."),
            "Inputs": "Status change commands (from kitchen staff or system), order identifier.",
            "Source": "Kitchen display actions, order cancellation module, and automated rules.",
            "Outputs": "Updated status field for the order; notifications to relevant clients.",
            "Destination": "Orders database, student order tracking view, admin views.",
            "Action": ("When a status update is triggered, the system verifies that the transition is allowed (e.g., cannot move "
                      "from completed back to preparing), updates the record, and pushes real-time updates via WebSocket or polling."),
            "Requirements": "Defined finite state machine for order statuses; real-time update mechanism.",
            "Pre-condition": "Order has been created and stored.",
            "Post-condition": "Order status reflects the latest valid state.",
            "Side effects": "Status transitions may write audit logs and affect analytics counters."
        }
    },
    {
        "title": "Order History and Detail Viewing",
        "rows": {
            "Function": "Provide students and admins with access to historical order information.",
            "Description": ("Allows users to view a list of their past orders and inspect detailed information such as items, "
                          "prices, pickup time, and payment method."),
            "Inputs": "User identifier, optional filter parameters (date range, status).",
            "Source": "Student dashboard, admin analytics or order-management screens.",
            "Outputs": "Paginated list of orders; detailed view for a selected order.",
            "Destination": "Order history page for students; admin back-office panels.",
            "Action": ("Upon request, the system queries the orders database by user and filter criteria, composes a list of "
                      "matching orders, and retrieves the full details when a user selects one."),
            "Requirements": "Indexed queries by user ID and date; secure access control.",
            "Pre-condition": "User is authenticated.",
            "Post-condition": "Requested historical data is displayed; no changes to underlying orders.",
            "Side effects": "Request may be logged for audit and performance tuning."
        }
    },
    {
        "title": "Order Cancellation",
        "rows": {
            "Function": "Allow students or admins to cancel orders under allowed conditions.",
            "Description": ("Supports cancellation of orders that are still in received state, ensuring that the kitchen is not "
                          "preparing the order and that any card payments are handled appropriately."),
            "Inputs": "Order identifier, cancellation request (student or admin), reason (optional).",
            "Source": "Student order detail page; admin order-management interface.",
            "Outputs": "Updated order status, optional refund trigger for paid orders, and confirmation message.",
            "Destination": "Orders database, student notification area, admin views.",
            "Action": ("When a cancellation request is received, the system checks the order status; if it is received, the status "
                      "is set to cancelled, any card payment is marked for refund or void, and the pickup slot capacity is released."),
            "Requirements": "Refund/void integration with payment provider; business rules for cancellation windows.",
            "Pre-condition": "Order exists and is in a cancellable state.",
            "Post-condition": "Order is marked as cancelled; associated resources (slots, analytics counts) are updated.",
            "Side effects": "Cancellation entries appear in audit logs and analytics."
        }
    },
    {
        "title": "Real-Time Order Notifications",
        "rows": {
            "Function": "Notify students of important order events in real time.",
            "Description": ("Delivers visual notifications when an order status changes (e.g., to preparing or ready) so that "
                          "students know when to arrive at the café."),
            "Inputs": "Order status change events, destination user session identifiers.",
            "Source": "Order status tracking module and WebSocket server.",
            "Outputs": "In-app notifications, badge updates, and refreshed order views.",
            "Destination": "Student web interface (orders page, notification area).",
            "Action": ("On every status change, the system pushes a message to connected clients subscribed to that order, "
                      "prompting them to refresh the visible state and show a toast or banner notification."),
            "Requirements": "WebSocket or similar push mechanism; subscription tracking per user.",
            "Pre-condition": "Order exists and user has an active session or will see the update on next refresh.",
            "Post-condition": "Student has up-to-date information about their order status.",
            "Side effects": "Notification events may be logged for reliability monitoring."
        }
    },
    {
        "title": "Kitchen Order Display and Queue Management",
        "rows": {
            "Function": "Provide kitchen staff with a real-time view of active orders organized by pickup time.",
            "Description": ("Shows all incoming and in-progress orders with key details such as items, quantities, pickup time, "
                          "and special instructions, enabling efficient preparation and fulfillment."),
            "Inputs": "Orders with statuses received, preparing, or ready.",
            "Source": "Orders database and order creation/status modules.",
            "Outputs": "Sorted list of orders with clear highlighting of approaching pickup times.",
            "Destination": "Kitchen display interface.",
            "Action": ("The system continuously retrieves or subscribes to order updates, presents them ordered by pickup time, "
                      "and visually distinguishes new or urgent orders. Staff actions on this screen update the order status."),
            "Requirements": "Stable connection between backend and kitchen device; role-based access for staff.",
            "Pre-condition": "Admin or kitchen staff user is authenticated.",
            "Post-condition": "Kitchen staff see an up-to-date queue and can act on each order.",
            "Side effects": "Status changes made here propagate to student views and analytics."
        }
    },
    {
        "title": "Admin Order Status Updates",
        "rows": {
            "Function": "Allow admins and kitchen staff to progress orders through the workflow.",
            "Description": ("Provides controls for changing an order's status from received → preparing → ready → completed "
                          "using the kitchen display or admin panel."),
            "Inputs": "Order identifier, selected new status.",
            "Source": "Kitchen display interface or admin order-management page.",
            "Outputs": "Updated status, visual feedback on the admin UI, and student notifications.",
            "Destination": "Orders database, student tracking view, audit logs.",
            "Action": ("When staff selects a new status, the system validates the transition, performs the update, and "
                      "broadcasts the change to subscribers."),
            "Requirements": "Role-based authorization; validation of allowed transitions.",
            "Pre-condition": "Order exists and user has the necessary privileges.",
            "Post-condition": "Order status reflects the new valid state.",
            "Side effects": "Status changes influence analytics such as completed orders and preparation times."
        }
    },
    {
        "title": "Menu Item Administration",
        "rows": {
            "Function": "Create, update, and delete cafeteria menu items.",
            "Description": ("Gives admins the ability to manage the menu catalogue including names in both languages, descriptions, "
                          "pricing, category, dietary tags, allergens, nutritional information, and images."),
            "Inputs": "Menu item fields (name, description, price, category, tags, image URL, etc.), item identifier for edits/deletes.",
            "Source": "Admin menu management interface.",
            "Outputs": "New or updated menu records; confirmation messages; refreshed menu lists.",
            "Destination": "Menu database; student-facing menu views.",
            "Action": ("On save, the system validates fields, writes changes to the database, and invalidates any relevant caches "
                      "so that students see the latest version. Deletes mark items as removed or archive them according to policy."),
            "Requirements": "Admin authentication; validation rules for required fields and value ranges.",
            "Pre-condition": "Admin user is logged in.",
            "Post-condition": "Menu catalogue reflects the latest changes.",
            "Side effects": "Changes are logged in the audit log and may impact analytics for item popularity."
        }
    },
    {
        "title": "Menu Availability and Special Pricing",
        "rows": {
            "Function": "Control real-time availability and promotional pricing of menu items.",
            "Description": ("Enables admins to toggle items as available/unavailable and configure special prices for daily "
                          "promotions or limited-time offers."),
            "Inputs": "Item identifier, availability flag, optional special price, special status flag.",
            "Source": "Admin menu management screen.",
            "Outputs": "Updated availability and pricing fields for each menu item.",
            "Destination": "Menu database, student menu views, order validation module.",
            "Action": ("When an admin updates availability or special pricing, the system saves the changes and immediately "
                      "propagates them to the menu display and order validation logic so that unavailable items cannot be added to new orders."),
            "Requirements": "Consistent price rules; immediate propagation to front-end cache or API layer.",
            "Pre-condition": "Menu item exists.",
            "Post-condition": "Students see correct availability and promotional prices in real time.",
            "Side effects": "Promotional periods may affect revenue analytics and are recorded for reporting."
        }
    },
    {
        "title": "Favorites Management",
        "rows": {
            "Function": "Allow students to save and manage favorite menu items.",
            "Description": ("Provides a quick-access list of frequently ordered items so students can reorder them with fewer steps."),
            "Inputs": "Student identifier, menu item identifier, add/remove actions.",
            "Source": "\"Add to favorites\" buttons on menu cards; favorites page actions.",
            "Outputs": "Updated favorites list for each student.",
            "Destination": "Favorites table in the database; favorites page in the student interface.",
            "Action": ("When a student marks or unmarks an item as favorite, the system updates the relation between the user "
                      "and the item and refreshes the favorites view accordingly."),
            "Requirements": "Unique user–item relationships; safeguards against duplicates.",
            "Pre-condition": "Student is authenticated.",
            "Post-condition": "Favorites list accurately reflects the student's preferences.",
            "Side effects": "Favorite counts can be used for popularity analytics."
        }
    },
    {
        "title": "Customer Feedback Submission",
        "rows": {
            "Function": "Collect structured feedback from students about their experience.",
            "Description": ("Allows students to submit feedback categorized by topic (e.g., food quality, service, menu suggestion, "
                          "general), optionally link it to a specific order, and provide a rating and comments."),
            "Inputs": "Student ID, category, rating, optional order ID, free-text message.",
            "Source": "Feedback form in the student interface.",
            "Outputs": "Feedback record stored in the feedback database; confirmation message.",
            "Destination": "Admin feedback management panel.",
            "Action": ("On submission, the system validates required fields, stores the feedback entry with status pending, "
                      "and acknowledges receipt to the student."),
            "Requirements": "Feedback storage schema; validation for categories and rating range.",
            "Pre-condition": "Student is authenticated.",
            "Post-condition": "Feedback entry exists and awaits review by admin.",
            "Side effects": "Feedback content may be used for service improvement reports."
        }
    },
    {
        "title": "Feedback Management (Admin)",
        "rows": {
            "Function": "Enable admins to review and manage student feedback.",
            "Description": ("Provides an interface to filter, read, and update the status of feedback items, and optionally "
                          "respond or record resolution notes."),
            "Inputs": "Feedback identifier, status updates, optional response or notes.",
            "Source": "Admin feedback management page.",
            "Outputs": "Updated feedback records with new statuses and notes.",
            "Destination": "Feedback database; optional student view of responses.",
            "Action": ("Admins filter feedback by category or status, open individual records, change the status to reviewed, "
                      "resolved, or dismissed, and save any comments; the system records these updates."),
            "Requirements": "Role-based admin access; audit logging for changes.",
            "Pre-condition": "Feedback entries exist.",
            "Post-condition": "Feedback statuses reflect current handling; unresolved items can be tracked.",
            "Side effects": "Statistics derived from feedback may appear in analytics."
        }
    },
    {
        "title": "Analytics and Reporting Dashboard",
        "rows": {
            "Function": "Provide summarized analytics and reports for café operations.",
            "Description": ("Generates visual dashboards showing sales, order counts, popular items, peak hours, and other "
                          "key performance metrics over configurable date ranges."),
            "Inputs": "Historical orders, revenue data, time filters, grouping parameters.",
            "Source": "Orders database and payment records.",
            "Outputs": "Charts, tables, and KPIs such as total revenue, total orders, average order value, and top-selling items.",
            "Destination": "Admin analytics page; CSV export files.",
            "Action": ("When admins select a date range or report type, the system aggregates relevant data, computes metrics, "
                      "renders charts, and prepares downloadable CSVs where applicable."),
            "Requirements": "Efficient aggregation queries; charting components; data export support.",
            "Pre-condition": "Sufficient historical data recorded.",
            "Post-condition": "Admins obtain up-to-date operational insights without altering raw data.",
            "Side effects": "Report generation load may influence database performance and should be monitored."
        }
    },
    {
        "title": "Language Toggle and Localization",
        "rows": {
            "Function": "Switch the user interface between English and Arabic with proper layout.",
            "Description": ("Allows users to choose their preferred language and ensures all labels, messages, and menu item "
                          "texts are displayed in that language, with right-to-left layout for Arabic."),
            "Inputs": "Language selection (EN or AR), translation dictionary entries, bilingual item fields.",
            "Source": "Language toggle control in the UI; translation resources.",
            "Outputs": "Localized interface text and properly oriented layout.",
            "Destination": "All visible pages and components in the web application.",
            "Action": ("When the user changes the language, the system applies the selected locale, reloads translation strings, "
                      "switches layout direction as needed, and persists the preference."),
            "Requirements": "Complete translation set; locale-aware UI components.",
            "Pre-condition": "User interface is loaded; translation bundles are available.",
            "Post-condition": "Application displays content consistently in the chosen language.",
            "Side effects": "Language choice stored for future sessions (e.g., in local storage)."
        }
    },
    {
        "title": "Theme and Currency Presentation",
        "rows": {
            "Function": "Manage visual theme (light/dark) and consistent currency display.",
            "Description": ("Provides toggling between light and dark modes and ensures that all prices are rendered using "
                          "the UAE Dirham symbol and correct formatting."),
            "Inputs": "Theme selection, currency format configuration, price values.",
            "Source": "User preference controls and system configuration.",
            "Outputs": "Updated UI theme and formatted price strings.",
            "Destination": "All pages where monetary values or theming are visible.",
            "Action": ("On theme change, the system updates global CSS or theme context; when prices are displayed, a formatting "
                      "helper adds the correct symbol and decimal precision."),
            "Requirements": "Persistent preference storage; unified currency formatting function.",
            "Pre-condition": "User interface loaded.",
            "Post-condition": "Theme selection and currency display remain consistent across pages and sessions.",
            "Side effects": "None beyond cosmetic changes."
        }
    },
    {
        "title": "Session Management and Security Logging",
        "rows": {
            "Function": "Handle user sessions and record important security-related events.",
            "Description": ("Maintains login sessions for students and admins, manages automatic timeouts, and logs key actions "
                          "such as logins, logouts, menu changes, and order status updates for accountability."),
            "Inputs": "User credentials, session tokens, action descriptors (e.g., \"menu item updated\").",
            "Source": "Authentication module, admin and student actions across the system.",
            "Outputs": "Active session records; audit log entries.",
            "Destination": "Sessions table; audit logs store.",
            "Action": ("On login, the system creates a session with expiration; periodic checks invalidate expired sessions. "
                      "For selected actions, the system writes an audit entry including user ID, timestamp, action type, and relevant identifiers."),
            "Requirements": "Secure token mechanisms; centralized logging strategy.",
            "Pre-condition": "User accounts and permissions are defined.",
            "Post-condition": "Valid sessions exist only for authenticated users; a trace of critical actions is preserved.",
            "Side effects": "Audit logs may be reviewed for troubleshooting or compliance."
        }
    }
]

# Process each requirement
for i, req in enumerate(requirements, 1):
    # Add heading for each requirement
    heading = document.add_heading(f'{i}. {req["title"]}', level=2)
    
    # Create table with 2 columns
    table = document.add_table(rows=len(req["rows"]), cols=2)
    table.style = 'Light Grid Accent 1'  # Use a nice table style
    
    # Populate table rows
    for idx, (key, value) in enumerate(req["rows"].items()):
        # Left column: Label (bold)
        cell_left = table.rows[idx].cells[0]
        cell_left.text = key
        cell_left.paragraphs[0].runs[0].bold = True
        
        # Right column: Value
        cell_right = table.rows[idx].cells[1]
        cell_right.text = value
    
    # Add spacing between requirements
    document.add_paragraph()

# Save the document
output_file = 'QuickDineFlow_Functional_Requirements.docx'
document.save(output_file)

print(f"Document successfully created: {output_file}")
print(f"Total requirements: {len(requirements)}")
print(f"You can now open the file in Microsoft Word!")

